"""
Enhanced Google Drive Service with multimodal support
Production-ready with file size limits, content validation, and duplicate detection
"""

import base64
import io
import logging
import os
from datetime import datetime
from typing import Any, Optional
from urllib.parse import urlencode

import aiohttp
import chardet
from PIL import Image

from app.services.memory_service import MemoryService

logger = logging.getLogger(__name__)

# File size limit: 250 MiB
MAX_FILE_SIZE = 250 * 1024 * 1024  # 250 MiB in bytes

# Supported content types for ingestion
SUPPORTED_MIME_TYPES = {
    # Text documents
    'text/plain': 'text',
    'text/html': 'text',
    'text/csv': 'text',
    'text/markdown': 'text',
    'application/json': 'text',
    'application/xml': 'text',

    # Code files (treated as text)
    'text/x-python': 'code',
    'text/x-java': 'code',
    'text/javascript': 'code',
    'text/x-c': 'code',
    'text/x-c++': 'code',
    'text/x-csharp': 'code',
    'text/x-ruby': 'code',
    'text/x-go': 'code',
    'text/x-rust': 'code',
    'text/x-typescript': 'code',

    # Documents
    'application/pdf': 'document',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'document',  # .docx
    'application/msword': 'document',  # .doc
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'spreadsheet',  # .xlsx
    'application/vnd.ms-excel': 'spreadsheet',  # .xls
    'application/epub+zip': 'ebook',

    # Google Docs formats
    'application/vnd.google-apps.document': 'gdoc',
    'application/vnd.google-apps.spreadsheet': 'gsheet',
    'application/vnd.google-apps.presentation': 'gslide',

    # Images
    'image/jpeg': 'image',
    'image/png': 'image',
    'image/gif': 'image',
    'image/webp': 'image',
    'image/svg+xml': 'image',
    'image/bmp': 'image',
    'image/tiff': 'image',
}


class GoogleDriveEnhanced:
    """Enhanced Google Drive integration with multimodal support"""

    def __init__(self):
        self.client_id = os.getenv("GOOGLE_CLIENT_ID")
        self.client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
        self.redirect_uri = os.getenv(
            "GOOGLE_REDIRECT_URI", "http://localhost:8001/api/v1/gdrive/callback"
        )
        self.memory_service = MemoryService()

        # Store tokens in memory for single user
        self.tokens = {}
        self.user_info = {}

        # Cache for duplicate detection
        self._ingested_files = set()  # Set of gdrive_ids that have been ingested

    def get_auth_url(self) -> str:
        """Generate Google OAuth URL"""
        params = {
            "client_id": self.client_id,
            "redirect_uri": self.redirect_uri,
            "response_type": "code",
            "scope": "https://www.googleapis.com/auth/drive.readonly email profile",
            "access_type": "offline",
            "prompt": "consent",
        }
        return f"https://accounts.google.com/o/oauth2/v2/auth?{urlencode(params)}"

    async def exchange_code(self, code: str) -> dict[str, Any]:
        """Exchange authorization code for tokens"""
        async with aiohttp.ClientSession() as session:
            data = {
                "code": code,
                "client_id": self.client_id,
                "client_secret": self.client_secret,
                "redirect_uri": self.redirect_uri,
                "grant_type": "authorization_code",
            }

            async with session.post("https://oauth2.googleapis.com/token", data=data) as resp:
                if resp.status == 200:
                    tokens = await resp.json()
                    self.tokens = tokens

                    # Get user info
                    await self.get_user_info(tokens["access_token"])

                    return {
                        "success": True,
                        "email": self.user_info.get("email"),
                        "tokens_stored": True,
                    }
                else:
                    error = await resp.text()
                    logger.error(f"Token exchange failed: {error}")
                    return {"success": False, "error": error}

    async def get_user_info(self, access_token: str) -> dict[str, Any]:
        """Get user info from Google"""
        async with aiohttp.ClientSession() as session:
            headers = {"Authorization": f"Bearer {access_token}"}
            async with session.get(
                "https://www.googleapis.com/oauth2/v2/userinfo", headers=headers
            ) as resp:
                if resp.status == 200:
                    self.user_info = await resp.json()
                    return self.user_info
                return {}

    async def list_files(self, folder_id: Optional[str] = None) -> list[dict[str, Any]]:
        """List files from Google Drive"""
        if not self.tokens.get("access_token"):
            return []

        async with aiohttp.ClientSession() as session:
            headers = {"Authorization": f'Bearer {self.tokens["access_token"]}'}

            # Build query
            query = "mimeType != 'application/vnd.google-apps.folder'"
            if folder_id:
                query = f"'{folder_id}' in parents and {query}"

            params = {
                "q": query,
                "fields": "files(id,name,mimeType,size,modifiedTime,webViewLink)",
                "pageSize": 100,
            }

            async with session.get(
                "https://www.googleapis.com/drive/v3/files", headers=headers, params=params
            ) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    files = data.get("files", [])
                    # Mark files that have already been ingested
                    for file in files:
                        file['already_ingested'] = file['id'] in self._ingested_files
                    return files
                else:
                    logger.error(f"Failed to list files: {resp.status}")
                    return []

    async def check_duplicate(self, file_id: str) -> bool:
        """
        Check if a file has already been ingested.
        Returns True if file exists in memories.
        """
        if file_id in self._ingested_files:
            return True

        # Check database for existing memory with this gdrive_id
        try:
            # Search for memories with this gdrive_id in metadata
            existing = await self.memory_service.search_memories(
                query=f"gdrive_id:{file_id}",
                limit=1
            )
            if existing:
                self._ingested_files.add(file_id)
                return True
        except Exception as e:
            logger.warning(f"Could not check for duplicate: {e}")

        return False

    def _validate_content_type(self, mime_type: str) -> tuple[bool, str]:
        """
        Validate if the content type is supported.
        Returns (is_supported, content_category)
        """
        # Check exact match first
        if mime_type in SUPPORTED_MIME_TYPES:
            return True, SUPPORTED_MIME_TYPES[mime_type]

        # Check for general categories
        if mime_type.startswith('text/'):
            return True, 'text'
        if mime_type.startswith('image/'):
            return True, 'image'
        if mime_type.startswith('application/') and 'json' in mime_type:
            return True, 'text'

        return False, 'unsupported'

    async def _process_image_content(self, content_bytes: bytes, metadata: dict[str, Any]) -> str:
        """
        Process image content - extract text via OCR and image metadata.
        """
        try:
            image = Image.open(io.BytesIO(content_bytes))

            # Extract image metadata
            img_info = {
                "format": image.format,
                "mode": image.mode,
                "size": f"{image.width}x{image.height}",
                "filename": metadata.get('name', 'unknown')
            }

            # Try OCR if pytesseract is available
            text_content = f"Image: {metadata.get('name', 'unknown')}\n"
            text_content += f"Format: {img_info['format']}\n"
            text_content += f"Dimensions: {img_info['size']}\n"
            text_content += f"Mode: {img_info['mode']}\n\n"

            try:
                import pytesseract
                # Extract text from image
                extracted_text = pytesseract.image_to_string(image)
                if extracted_text.strip():
                    text_content += "Extracted Text:\n" + extracted_text
            except ImportError:
                text_content += "[OCR not available - install pytesseract for text extraction]"
            except Exception as e:
                logger.warning(f"OCR failed: {e}")
                text_content += "[OCR extraction failed]"

            return text_content

        except Exception as e:
            logger.error(f"Failed to process image: {e}")
            return f"[Image processing failed: {e}]"

    async def _process_pdf_content(self, content_bytes: bytes, metadata: dict[str, Any]) -> str:
        """
        Process PDF content - extract text and metadata.
        """
        try:
            import pypdf

            pdf_reader = pypdf.PdfReader(io.BytesIO(content_bytes))

            text_content = f"PDF Document: {metadata.get('name', 'unknown')}\n"
            text_content += f"Pages: {len(pdf_reader.pages)}\n\n"

            # Extract text from all pages
            for i, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content += f"--- Page {i} ---\n{text}\n\n"

            # Extract metadata if available
            if pdf_reader.metadata:
                text_content += "\nDocument Metadata:\n"
                for key, value in pdf_reader.metadata.items():
                    text_content += f"  {key}: {value}\n"

            return text_content

        except Exception as e:
            logger.error(f"Failed to process PDF: {e}")
            return f"[PDF processing failed: {e}]"

    async def _process_docx_content(self, content_bytes: bytes, metadata: dict[str, Any]) -> str:
        """
        Process DOCX content - extract text and metadata.
        """
        try:
            import docx

            doc = docx.Document(io.BytesIO(content_bytes))

            text_content = f"Word Document: {metadata.get('name', 'unknown')}\n\n"

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"

            # Extract tables
            if doc.tables:
                text_content += "\n\nTables Found:\n"
                for i, table in enumerate(doc.tables, 1):
                    text_content += f"\nTable {i}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_content += row_text + "\n"

            return text_content

        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            return f"[DOCX processing failed: {e}]"

    async def _fetch_file_metadata(self, session: aiohttp.ClientSession, file_id: str, headers: dict) -> Optional[dict]:
        """Fetch file metadata from Google Drive."""
        metadata_url = f"https://www.googleapis.com/drive/v3/files/{file_id}"
        params = {"fields": "id,name,mimeType,modifiedTime,webViewLink,size"}
        async with session.get(metadata_url, headers=headers, params=params) as resp:
            if resp.status != 200:
                logger.error(f"Failed to get metadata for file {file_id}: {resp.status}")
                return None
            return await resp.json()

    async def _download_file_content(self, session: aiohttp.ClientSession, file_id: str,
                                    mime_type: str, headers: dict) -> Optional[bytes]:
        """Download file content from Google Drive."""
        # Handle Google Docs, Sheets, Slides by exporting them
        if "google-apps" in mime_type:
            export_mime = "text/plain"
            if "spreadsheet" in mime_type:
                export_mime = "text/csv"
            elif "presentation" in mime_type:
                export_mime = "text/plain"
            elif "document" in mime_type:
                export_mime = "text/plain"

            url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export"
            params = {"mimeType": export_mime}
        else:
            # Download other files directly
            url = f"https://www.googleapis.com/drive/v3/files/{file_id}"
            params = {"alt": "media"}

        async with session.get(url, headers=headers, params=params) as resp:
            if resp.status == 200:
                return await resp.read()
            else:
                logger.error(f"Failed to download file content: {resp.status}")
                return None

    async def _process_file_content(self, content_bytes: bytes, mime_type: str,
                                   content_category: str, metadata: dict) -> Optional[str]:
        """Process file content based on type."""
        if content_category == 'image':
            return await self._process_image_content(content_bytes, metadata)
        elif mime_type == 'application/pdf':
            return await self._process_pdf_content(content_bytes, metadata)
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          'application/msword']:
            return await self._process_docx_content(content_bytes, metadata)
        else:
            # Text-based content - detect encoding
            detected = chardet.detect(content_bytes)
            encoding = detected.get('encoding', 'utf-8') if detected else 'utf-8'

            try:
                return content_bytes.decode(encoding)
            except UnicodeDecodeError:
                # Fallback to utf-8 with errors ignored
                logger.warning("Had to ignore encoding errors for file")
                return content_bytes.decode('utf-8', errors='ignore')

    async def _create_memory_from_file(self, content: str, metadata: dict, content_category: str,
                                      content_bytes: Optional[bytes], file_size: Optional[str]) -> Optional[dict]:
        """Create memory from processed file content."""
        # Prepare metadata for the memory
        memory_metadata = {
            "source": "google_drive",
            "gdrive_id": metadata.get("id"),
            "gdrive_mimetype": metadata.get("mimeType"),
            "gdrive_weblink": metadata.get("webViewLink"),
            "gdrive_modified_time": metadata.get("modifiedTime"),
            "original_filename": metadata.get("name"),
            "content_category": content_category,
            "file_size": file_size,
        }

        # Add image data if we have it
        if content_bytes and content_category == 'image':
            # Store base64 encoded thumbnail for later reference
            try:
                img = Image.open(io.BytesIO(content_bytes))
                # Create thumbnail (max 256x256)
                img.thumbnail((256, 256), Image.Resampling.LANCZOS)
                thumb_io = io.BytesIO()
                img.save(thumb_io, format='PNG')
                thumb_b64 = base64.b64encode(thumb_io.getvalue()).decode('utf-8')
                memory_metadata['thumbnail'] = f"data:image/png;base64,{thumb_b64}"
            except Exception as e:
                logger.warning(f"Could not create thumbnail: {e}")

        try:
            # Adjust importance based on content type
            importance_scores = {
                'code': 0.8,
                'document': 0.75,
                'gdoc': 0.75,
                'gsheet': 0.7,
                'image': 0.6,
                'text': 0.7,
                'ebook': 0.65,
                'spreadsheet': 0.7,
            }
            importance = importance_scores.get(content_category, 0.7)

            created_memory = await self.memory_service.create_memory(
                content=content,
                metadata=memory_metadata,
                tags=["gdrive", content_category, metadata.get("mimeType")],
                importance_score=importance
            )

            return created_memory
        except Exception as e:
            logger.error(f"Failed to create memory: {e}")
            return None

    async def ingest_file(self, file_id: str, force: bool = False) -> Optional[dict[str, Any]]:
        """
        Reads a file from Google Drive, sends it for embedding, and stores it as a memory.

        Args:
            file_id: Google Drive file ID
            force: If True, skip duplicate check and re-ingest the file

        Returns:
            Created memory dict or None if failed
        """
        if not self.tokens.get("access_token"):
            logger.error("Not authenticated with Google Drive.")
            return None

        # Check for duplicate unless force is True
        if not force and await self.check_duplicate(file_id):
            logger.info(f"File {file_id} already ingested, skipping.")
            return {"status": "skipped", "reason": "duplicate", "file_id": file_id}

        async with aiohttp.ClientSession() as session:
            headers = {"Authorization": f'Bearer {self.tokens["access_token"]}'}

            # 1. Get file metadata
            metadata = await self._fetch_file_metadata(session, file_id, headers)
            if not metadata:
                return None

            # 2. Validate content type
            mime_type = metadata.get("mimeType", "")
            is_supported, content_category = self._validate_content_type(mime_type)

            if not is_supported:
                logger.warning(f"Unsupported file type: {mime_type} for file {metadata.get('name')}")
                return {"status": "skipped", "reason": "unsupported_type",
                        "mime_type": mime_type, "file_name": metadata.get('name')}

            # 3. Check file size (if available)
            file_size = metadata.get('size')
            if file_size:
                file_size_int = int(file_size)
                if file_size_int > MAX_FILE_SIZE:
                    logger.warning(f"File too large: {file_size_int} bytes (max: {MAX_FILE_SIZE})")
                    return {"status": "skipped", "reason": "file_too_large",
                            "size": file_size_int, "max_size": MAX_FILE_SIZE}

            # 4. Download file content
            content_bytes = await self._download_file_content(session, file_id, mime_type, headers)
            if not content_bytes:
                return None

            # Check size after download
            if len(content_bytes) > MAX_FILE_SIZE:
                logger.warning(f"Downloaded file too large: {len(content_bytes)} bytes")
                return {"status": "skipped", "reason": "file_too_large",
                        "size": len(content_bytes), "max_size": MAX_FILE_SIZE}

            # 5. Process content based on type
            content = await self._process_file_content(content_bytes, mime_type, content_category, metadata)
            if not content:
                logger.error(f"File {file_id} has no content.")
                return None

            # 6. Create memory
            logger.info(f"Ingesting file: {metadata.get('name')} (type: {content_category})")

            created_memory = await self._create_memory_from_file(
                content, metadata, content_category,
                content_bytes if content_category == 'image' else None,
                file_size
            )

            if created_memory:
                # Mark as ingested
                self._ingested_files.add(file_id)
                logger.info(f"Successfully created memory {created_memory['id']} for file {metadata.get('name')}")
                return created_memory

            return None

    def is_connected(self) -> bool:
        """Check if we have valid tokens"""
        return bool(self.tokens.get("access_token"))

    def get_connection_status(self) -> dict[str, Any]:
        """Get current connection status"""
        if self.is_connected():
            return {
                "connected": True,
                "user_email": self.user_info.get("email"),
                "user_name": self.user_info.get("name"),
                "last_checked": datetime.utcnow().isoformat(),
                "ingested_files_count": len(self._ingested_files),
                "max_file_size_mb": MAX_FILE_SIZE / (1024 * 1024),
                "supported_types": list(SUPPORTED_MIME_TYPES.keys())[:10]  # Sample
            }
        else:
            return {"connected": False, "error": "Not authenticated", "requires_auth": True}

    async def batch_ingest(self, file_ids: list[str], progress_callback=None) -> dict[str, Any]:
        """
        Batch ingest multiple files with progress tracking.

        Args:
            file_ids: List of Google Drive file IDs
            progress_callback: Optional async callback for progress updates

        Returns:
            Summary of ingestion results
        """
        results = {
            "total": len(file_ids),
            "successful": 0,
            "skipped": 0,
            "failed": 0,
            "details": []
        }

        for i, file_id in enumerate(file_ids, 1):
            try:
                result = await self.ingest_file(file_id)

                if result:
                    if result.get("status") == "skipped":
                        results["skipped"] += 1
                        results["details"].append({
                            "file_id": file_id,
                            "status": "skipped",
                            "reason": result.get("reason")
                        })
                    else:
                        results["successful"] += 1
                        results["details"].append({
                            "file_id": file_id,
                            "status": "success",
                            "memory_id": result.get("id")
                        })
                else:
                    results["failed"] += 1
                    results["details"].append({
                        "file_id": file_id,
                        "status": "failed"
                    })

                # Progress callback
                if progress_callback:
                    await progress_callback({
                        "current": i,
                        "total": len(file_ids),
                        "percentage": (i / len(file_ids)) * 100,
                        "file_id": file_id
                    })

            except Exception as e:
                logger.error(f"Error ingesting file {file_id}: {e}")
                results["failed"] += 1
                results["details"].append({
                    "file_id": file_id,
                    "status": "error",
                    "error": str(e)
                })

        return results


# Global instance for single user
google_drive_enhanced = GoogleDriveEnhanced()

