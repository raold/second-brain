"""
File parsers for various document types
"""

from pathlib import Path
from typing import Any

import aiofiles

# Document parsing libraries
try:
    import PyPDF2
    from pdfplumber import PDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import html2text
    from bs4 import BeautifulSoup
    HAS_HTML = True
except ImportError:
    HAS_HTML = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import openpyxl
    import pandas as pd
    HAS_SPREADSHEET = True
except ImportError:
    HAS_SPREADSHEET = False

from app.ingestion.engine import FileParser
from app.utils.logger import get_logger
from typing import List
from typing import Any

logger = get_logger(__name__)


class PDFParser(FileParser):
    """Parser for PDF files"""

    SUPPORTED_TYPES = {'application/pdf'}

    def __init__(self):
        if not HAS_PDF:
            raise ImportError(
                "PDF parsing libraries not installed. "
                "Install with: pip install PyPDF2 pdfplumber"
            )

    async def parse(self, file_path: Path) -> dict[str, Any]:
        """Parse PDF file and extract text"""
        content_parts = []
        metadata = {
            'format': 'pdf',
            'pages': 0,
            'has_images': False,
            'has_tables': False
        }

        try:
            # Try pdfplumber first for better text extraction
            with PDF.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"[Page {page_num}]\n{text}")

                    # Check for tables
                    tables = page.extract_tables()
                    if tables:
                        metadata['has_tables'] = True
                        for table in tables:
                            # Convert table to text representation
                            table_text = self._table_to_text(table)
                            content_parts.append(f"[Table on Page {page_num}]\n{table_text}")

                    # Check for images
                    if page.images:
                        metadata['has_images'] = True

        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")

            # Fallback to PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata['pages'] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"[Page {page_num}]\n{text}")

        return {
            'content': '\n\n'.join(content_parts),
            'metadata': metadata
        }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    @staticmethod
    def _table_to_text(table: list[list[str]]) -> str:
        """Convert table data to formatted text"""
        if not table:
            return ""

        # Calculate column widths
        col_widths = []
        for col_idx in range(len(table[0])):
            max_width = max(
                len(str(row[col_idx]) if col_idx < len(row) else "")
                for row in table
            )
            col_widths.append(max_width)

        # Format table
        lines = []
        for row in table:
            formatted_row = " | ".join(
                str(cell).ljust(col_widths[i])
                for i, cell in enumerate(row)
            )
            lines.append(formatted_row)

        return "\n".join(lines)


class DocxParser(FileParser):
    """Parser for DOCX files"""

    SUPPORTED_TYPES = {
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    }

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError(
                "DOCX parsing library not installed. "
                "Install with: pip install python-docx"
            )

    async def parse(self, file_path: Path) -> dict[str, Any]:
        """Parse DOCX file and extract text"""
        doc = Document(file_path)
        content_parts = []
        metadata = {
            'format': 'docx',
            'paragraphs': 0,
            'tables': 0,
            'lists': 0,
            'images': 0
        }

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)
                metadata['paragraphs'] += 1

                # Check if it's a list item
                if para.style.name.startswith('List'):
                    metadata['lists'] += 1

        # Extract tables
        for table in doc.tables:
            metadata['tables'] += 1
            table_text = self._extract_table_text(table)
            content_parts.append(f"[Table]\n{table_text}")

        # Count images (relationships)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                metadata['images'] += 1

        return {
            'content': '\n\n'.join(content_parts),
            'metadata': metadata
        }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES

    @staticmethod
    def _extract_table_text(table) -> str:
        """Extract text from DOCX table"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


class HTMLParser(FileParser):
    """Parser for HTML files"""

    SUPPORTED_TYPES = {'text/html', 'application/xhtml+xml'}

    def __init__(self):
        if not HAS_HTML:
            raise ImportError(
                "HTML parsing libraries not installed. "
                "Install with: pip install beautifulsoup4 html2text"
            )
        self.h2t = html2text.HTML2Text()
        self.h2t.ignore_links = False
        self.h2t.ignore_images = False
        self.h2t.body_width = 0  # Don't wrap lines

    async def parse(self, file_path: Path) -> dict[str, Any]:
        """Parse HTML file and extract text"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = await f.read()

        # Parse with BeautifulSoup for metadata
        soup = BeautifulSoup(html_content, 'html.parser')

        metadata = {
            'format': 'html',
            'title': '',
            'links': [],
            'images': [],
            'headers': {}
        }

        # Extract title
        title_tag = soup.find('title')
        if title_tag:
            metadata['title'] = title_tag.get_text(strip=True)

        # Extract links
        for link in soup.find_all('a', href=True):
            metadata['links'].append({
                'text': link.get_text(strip=True),
                'href': link['href']
            })

        # Extract images
        for img in soup.find_all('img', src=True):
            metadata['images'].append({
                'src': img['src'],
                'alt': img.get('alt', '')
            })

        # Count headers
        for i in range(1, 7):
            headers = soup.find_all(f'h{i}')
            if headers:
                metadata['headers'][f'h{i}'] = len(headers)

        # Convert to markdown for better text representation
        markdown_content = self.h2t.handle(html_content)

        return {
            'content': markdown_content,
            'metadata': metadata
        }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES


class ImageParser(FileParser):
    """Parser for image files with OCR"""

    SUPPORTED_TYPES = {
        'image/jpeg',
        'image/jpg',
        'image/png',
        'image/gif',
        'image/bmp',
        'image/tiff'
    }

    def __init__(self):
        if not HAS_OCR:
            raise ImportError(
                "OCR libraries not installed. "
                "Install with: pip install pytesseract pillow"
            )

    async def parse(self, file_path: Path) -> dict[str, Any]:
        """Parse image file and extract text using OCR"""
        # Open image
        image = Image.open(file_path)

        metadata = {
            'format': image.format,
            'size': image.size,
            'mode': image.mode,
            'has_text': False
        }

        try:
            # Perform OCR
            text = pytesseract.image_to_string(image)

            if text.strip():
                metadata['has_text'] = True
                content = text
            else:
                content = f"[Image: {file_path.name} - No text detected]"

            # Get additional OCR data
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidence_scores = [
                conf for conf in ocr_data['conf']
                if isinstance(conf, int) and conf > 0
            ]

            if confidence_scores:
                metadata['ocr_confidence'] = sum(confidence_scores) / len(confidence_scores)

        except Exception as e:
            logger.error(f"OCR failed: {e}")
            content = f"[Image: {file_path.name} - OCR failed]"

        return {
            'content': content,
            'metadata': metadata
        }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES


class SpreadsheetParser(FileParser):
    """Parser for spreadsheet files"""

    SUPPORTED_TYPES = {
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',  # XLSX
        'application/vnd.ms-excel',  # XLS
        'text/csv',
        'application/csv'
    }

    def __init__(self):
        if not HAS_SPREADSHEET:
            raise ImportError(
                "Spreadsheet parsing libraries not installed. "
                "Install with: pip install pandas openpyxl"
            )

    async def parse(self, file_path: Path) -> dict[str, Any]:
        """Parse spreadsheet file and extract data"""
        content_parts = []
        metadata = {
            'format': 'spreadsheet',
            'sheets': [],
            'total_rows': 0,
            'total_columns': 0
        }

        try:
            if file_path.suffix.lower() == '.csv':
                # Handle CSV files
                df = pd.read_csv(file_path)
                sheet_name = 'Sheet1'

                metadata['sheets'].append({
                    'name': sheet_name,
                    'rows': len(df),
                    'columns': len(df.columns)
                })
                metadata['total_rows'] = len(df)
                metadata['total_columns'] = len(df.columns)

                # Convert to text
                content_parts.append(f"[{sheet_name}]")
                content_parts.append(df.to_string())

            else:
                # Handle Excel files
                excel_file = pd.ExcelFile(file_path)

                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)

                    metadata['sheets'].append({
                        'name': sheet_name,
                        'rows': len(df),
                        'columns': len(df.columns)
                    })
                    metadata['total_rows'] += len(df)
                    metadata['total_columns'] = max(
                        metadata['total_columns'],
                        len(df.columns)
                    )

                    # Convert to text
                    content_parts.append(f"[Sheet: {sheet_name}]")
                    content_parts.append(df.to_string())
                    content_parts.append("")  # Empty line between sheets

        except Exception as e:
            logger.error(f"Failed to parse spreadsheet: {e}")
            # Try simpler CSV parsing
            if file_path.suffix.lower() == '.csv':
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    content = await f.read()
                content_parts.append(content)

        return {
            'content': '\n\n'.join(content_parts),
            'metadata': metadata
        }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES


class MarkdownParser(FileParser):
    """Parser for Markdown files"""

    SUPPORTED_TYPES = {
        'text/markdown',
        'text/x-markdown',
        'application/x-markdown'
    }

    async def parse(self, file_path: Path) -> dict[str, Any]:
        """Parse Markdown file"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()

        # Extract metadata from frontmatter if present
        metadata = {'format': 'markdown'}

        if content.startswith('---'):
            try:
                import yaml
                parts = content.split('---', 2)
                if len(parts) >= 3:
                    frontmatter = yaml.safe_load(parts[1])
                    metadata['frontmatter'] = frontmatter
                    content = parts[2]
            except:
                pass

        # Count various elements
        lines = content.split('\n')
        metadata['headers'] = {}
        metadata['links'] = 0
        metadata['images'] = 0
        metadata['code_blocks'] = 0

        for line in lines:
            # Count headers
            if line.startswith('#'):
                level = len(line.split()[0])
                header_key = f'h{level}'
                metadata['headers'][header_key] = metadata['headers'].get(header_key, 0) + 1

            # Count links
            metadata['links'] += line.count('](')

            # Count images
            metadata['images'] += line.count('![')

            # Count code blocks
            if line.strip() == '```':
                metadata['code_blocks'] += 0.5  # Will sum to integer

        metadata['code_blocks'] = int(metadata['code_blocks'])

        return {
            'content': content,
            'metadata': metadata
        }

    def supports(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_TYPES or mime_type == 'text/plain'


# Export all parsers
__all__ = [
    'PDFParser',
    'DocxParser',
    'HTMLParser',
    'ImageParser',
    'SpreadsheetParser',
    'MarkdownParser'
]
